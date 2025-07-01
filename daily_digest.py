# Daily Digest Generator
# Consolidated newsletter generation for EquiLend D&A

import os
import pandas as pd
import numpy as np
import datetime as dt
import google.generativeai as genai
from typing import List, Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor
from jinja2 import Environment, FileSystemLoader

# EquiLend color palette
BLUE = RGBColor(4, 141, 210)
ORANGE = RGBColor(245, 118, 0)
BLACK = RGBColor(0, 0, 0)
PURPLE = RGBColor(105, 1, 208)
GREEN = RGBColor(29, 183, 26)
RED = RGBColor(255, 0, 0)

TODAY = dt.date.today()

class DailyDigestGenerator:
    """Main class for generating EquiLend Daily Digest"""
    
    def __init__(self):
        self.gemini_key = os.getenv("GEMINI_API_KEY")
        if self.gemini_key:
            genai.configure(api_key=self.gemini_key)
    
    def load_daily_data(self) -> pd.DataFrame:
        """Load daily factor data from SQL or CSV"""
        if os.getenv("USE_SQL") == "1":
            try:
                import sqlalchemy
                engine = sqlalchemy.create_engine(os.getenv("SQL_URI"))
                query = "SELECT * FROM equilend.factors_daily WHERE ds = current_date"
                return pd.read_sql(query, engine)
            except Exception as e:
                print(f"Error loading data from SQL: {e}")
        
        # Fallback to CSV
        csv_path = f"/data/daily_factors_{TODAY}.csv"
        if os.path.exists(csv_path):
            return pd.read_csv(csv_path)
        
        return pd.DataFrame()  # Return empty dataframe if no data
    
    def extract_key_bullets(self, df: pd.DataFrame) -> List[str]:
        """Extract key data bullets from factor dataframe"""
        bullets = []
        
        if df.empty:
            return ["No data available for today"]
        
        try:
            # Top loan quantity changes
            if 'On Loan Quantity Month Diff' in df.columns:
                top_loan = df.nlargest(1, 'On Loan Quantity Month Diff').iloc[0]
                bullets.append(f"{top_loan.get('ticker', 'Unknown')}: On-loan quantity up {top_loan['On Loan Quantity Month Diff']:.1f}% MoM")
            
            # Top utilization
            if 'Active Utilization (%)' in df.columns:
                top_util = df.nlargest(1, 'Active Utilization (%)').iloc[0]
                bullets.append(f"{top_util.get('ticker', 'Unknown')}: Utilization at {top_util['Active Utilization (%)']:.1f}%")
            
            # Top short squeeze scores
            if 'SSR_v4' in df.columns:
                top_squeeze = df.nlargest(1, 'SSR_v4').iloc[0]
                bullets.append(f"{top_squeeze.get('ticker', 'Unknown')}: Enhanced squeeze score {top_squeeze['SSR_v4']:.2f}")
        
        except Exception:
            bullets = ["Data processing error - manual review required"]
        
        return bullets[:3]  # Limit to 3 bullets
    
    def generate_takeaways(self, bullets: List[str]) -> List[str]:
        """Generate three key takeaways using Gemini"""
        if not self.gemini_key or not bullets:
            return [
                "Market data shows continued securities lending activity",
                "Borrow costs remain elevated in select names",
                "Utilization patterns suggest ongoing supply-demand imbalances"
            ]
        
        try:
            model = genai.GenerativeModel('gemini-pro')
            
            prompt = f"""Style: EquiLend Data & Analytics. Tone: even, data-forward.
            Provide exactly three bullet takeaways covering macro themes visible
            in these data points: {bullets}. Do NOT add dramatic language."""
            
            response = model.generate_content(prompt)
            
            takeaways = [line.strip('- ') for line in response.text.split('\n') if line.strip()]
            return takeaways[:3]
        
        except Exception:
            return [
                "Market data analysis in progress",
                "Securities lending metrics under review", 
                "Factor analysis continues to evolve"
            ]
    
    def build_word_document(self, avg_fee: float, avg_fee_chg: float, 
                           avg_util: float, avg_util_chg: float,
                           headlines: List[Dict[str, str]], 
                           specials_df: pd.DataFrame,
                           takeaways: List[str],
                           output_filename: str) -> None:
        """Build Word document with EquiLend branding"""
        
        try:
            doc = Document('EquiLend Letterhead_Please Download.docx')
        except Exception:
            doc = Document()
        
        # Title
        title_run = doc.add_paragraph().add_run('EquiLend D&A Daily Digest')
        title_run.font.color.rgb = BLUE
        title_run.font.size = Pt(20)
        title_run.bold = True
        
        # Date
        date_run = doc.add_paragraph().add_run(TODAY.strftime('%B %d, %Y'))
        date_run.font.color.rgb = ORANGE
        
        # Market Notes
        market_run = doc.add_paragraph().add_run('📈 MARKET NOTES:')
        market_run.font.color.rgb = BLUE
        market_run.bold = True
        
        # Fee metrics
        fee_color = GREEN if avg_fee_chg > 0 else RED
        fee_arrow = '↑' if avg_fee_chg > 0 else '↓'
        fee_run = doc.add_paragraph().add_run(f'Fee All: {avg_fee:.2f}bps {fee_arrow}{abs(avg_fee_chg):.2f}%')
        fee_run.font.color.rgb = fee_color
        
        # Utilization metrics
        util_color = GREEN if avg_util_chg > 0 else RED
        util_arrow = '↑' if avg_util_chg > 0 else '↓'
        util_run = doc.add_paragraph().add_run(f'Utilization: {avg_util:.2f}% {util_arrow}{abs(avg_util_chg):.2f}%')
        util_run.font.color.rgb = util_color
        
        # Headlines
        headline_run = doc.add_paragraph().add_run('🔥 Major Headlines & What Our Data Shows')
        headline_run.font.color.rgb = BLUE
        headline_run.bold = True
        
        for headline in headlines:
            title_run = doc.add_paragraph().add_run(headline['title'])
            title_run.font.color.rgb = BLACK
            title_run.bold = True
            
            data_run = doc.add_paragraph().add_run('Data: ' + headline['bullet'])
            data_run.font.color.rgb = PURPLE
            data_run.bold = True
        
        # Specials table
        if not specials_df.empty:
            table = doc.add_table(rows=1, cols=len(specials_df.columns))
            
            # Header row
            for i, col in enumerate(specials_df.columns):
                cell_run = table.rows[0].cells[i].paragraphs[0].add_run(col)
                cell_run.font.color.rgb = BLUE
                cell_run.bold = True
            
            # Data rows
            for _, row in specials_df.iterrows():
                cells = table.add_row().cells
                for j, val in enumerate(row):
                    cells[j].text = str(val)
        
        # Takeaways
        takeaway_run = doc.add_paragraph().add_run('💡 KEY TAKEAWAYS')
        takeaway_run.font.color.rgb = BLUE
        takeaway_run.bold = True
        
        for takeaway in takeaways:
            bullet_run = doc.add_paragraph().add_run('• ' + takeaway)
            bullet_run.font.color.rgb = BLACK
        
        doc.save(output_filename)
    
    def generate_daily_digest(self, manual_inputs: Dict[str, Any]) -> Dict[str, Any]:
        """Main method to generate complete daily digest"""
        
        # Load data
        df = self.load_daily_data()
        
        # Extract insights
        bullets = self.extract_key_bullets(df)
        takeaways = self.generate_takeaways(bullets)
        
        # Create specials dataframe
        if not df.empty:
            try:
                top_indices = df.nlargest(5, 'Fee All (BPS)').index if 'Fee All (BPS)' in df.columns else df.head(5).index
                specials_df = df.loc[top_indices, ['ticker', 'sec_desc', 'industry']].copy()
                specials_df.columns = ['Ticker', 'Security Description', 'Industry']
            except Exception:
                specials_df = pd.DataFrame(columns=['Ticker', 'Security Description', 'Industry'])
        else:
            specials_df = pd.DataFrame(columns=['Ticker', 'Security Description', 'Industry'])
        
        # Build headlines
        headlines = [
            {'title': manual_inputs.get('headline_1', 'Market Update'), 'bullet': bullets[0] if bullets else 'No data'},
            {'title': manual_inputs.get('headline_2', 'Sector Focus'), 'bullet': bullets[1] if len(bullets) > 1 else 'No data'},
            {'title': manual_inputs.get('headline_3', 'Technical Analysis'), 'bullet': bullets[2] if len(bullets) > 2 else 'No data'}
        ]
        
        # Generate Word document
        output_file = f'Daily_Digest_{TODAY}.docx'
        self.build_word_document(
            manual_inputs.get('avg_fee', 45.0),
            manual_inputs.get('avg_fee_chg', 0.0),
            manual_inputs.get('avg_util', 7.0),
            manual_inputs.get('avg_util_chg', 0.0),
            headlines,
            specials_df,
            takeaways,
            output_file
        )
        
        return {
            'output_file': output_file,
            'bullets': bullets,
            'takeaways': takeaways,
            'specials': specials_df,
            'headlines': headlines,
            'data_rows': len(df)
        }

# Convenience function for quick generation
def generate_digest(avg_fee=45.0, avg_fee_chg=0.0, avg_util=7.0, avg_util_chg=0.0,
                   headline_1="Market Update", headline_2="Sector Focus", headline_3="Technical Analysis"):
    """Quick function to generate daily digest with manual inputs"""
    
    generator = DailyDigestGenerator()
    manual_inputs = {
        'avg_fee': avg_fee,
        'avg_fee_chg': avg_fee_chg,
        'avg_util': avg_util,
        'avg_util_chg': avg_util_chg,
        'headline_1': headline_1,
        'headline_2': headline_2,
        'headline_3': headline_3
    }
    
    return generator.generate_daily_digest(manual_inputs)
